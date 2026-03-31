"""DOCX document parser.

Uses python-docx to extract text, sections (based on paragraph styles),
tables, and footnotes from Microsoft Word documents.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

import structlog

from paperverifier.models.document import (
    ParsedDocument,
    Section,
)
from paperverifier.parsers.base import BaseParser
from paperverifier.security.input_validator import (
    InputValidationError,
    validate_file_path,
    validate_uploaded_file,
)

logger = structlog.get_logger(__name__)

_MAX_DOCX_SIZE = 100 * 1024 * 1024  # 100 MB

# Heading style names used by python-docx.
_HEADING_STYLES = {
    "Heading 1": 1,
    "Heading 2": 2,
    "Heading 3": 3,
    "Heading 4": 4,
    "Heading 5": 5,
    "Heading 6": 6,
    "Title": 0,
    "Subtitle": 0,
}


class DOCXParser(BaseParser):
    """Parse DOCX (Office Open XML) documents into :class:`ParsedDocument`.

    Maps Word paragraph styles (``Heading 1``, ``Heading 2``, etc.) to
    document sections.  Tables and footnotes are extracted and appended
    to the relevant section text.
    """

    async def parse(self, source: str | bytes, **kwargs: object) -> ParsedDocument:
        """Parse a DOCX from a file path or raw bytes.

        Args:
            source: Either a file path (str) or raw DOCX bytes.
            **kwargs: Optional ``allowed_dir`` (:class:`Path`).

        Returns:
            A fully populated :class:`ParsedDocument`.
        """
        try:
            import docx  # type: ignore[import-untyped]
        except ImportError:
            raise RuntimeError(
                "python-docx is required for DOCX parsing. "
                "Install it with: pip install python-docx"
            )

        source_path = ""

        if isinstance(source, bytes):
            if len(source) == 0:
                raise InputValidationError("DOCX file is empty.")
            if len(source) > _MAX_DOCX_SIZE:
                raise InputValidationError(
                    f"DOCX file too large: {len(source):,} bytes exceeds "
                    f"maximum of {_MAX_DOCX_SIZE:,} bytes."
                )
            # python-docx can open from a file-like object.
            import io

            doc = docx.Document(io.BytesIO(source))
        else:
            path = Path(source)
            allowed_dir = kwargs.get("allowed_dir")
            if allowed_dir is not None:
                validate_file_path(path, Path(allowed_dir))  # type: ignore[arg-type]

            if not path.exists():
                raise InputValidationError(f"DOCX file not found: {source}")
            if path.stat().st_size > _MAX_DOCX_SIZE:
                raise InputValidationError(
                    f"DOCX file too large: {path.stat().st_size:,} bytes exceeds "
                    f"maximum of {_MAX_DOCX_SIZE:,} bytes."
                )
            source_path = str(path)
            doc = docx.Document(str(path))

        # Extract core properties for metadata.
        metadata: dict[str, Any] = {}
        title: str | None = None
        authors: list[str] = []

        try:
            props = doc.core_properties
            if props.title:
                title = props.title
                metadata["title"] = title
            if props.author:
                authors = [a.strip() for a in props.author.split(",") if a.strip()]
                metadata["author"] = props.author
            if props.subject:
                metadata["subject"] = props.subject
            if props.keywords:
                metadata["keywords"] = props.keywords
        except Exception:
            pass

        # Walk paragraphs and build section structure.
        sections, full_text = self._extract_sections(doc)

        # Extract tables and append to metadata.
        tables_text = self._extract_tables(doc)
        if tables_text:
            metadata["tables"] = tables_text
            full_text += "\n\n" + tables_text

        # Extract footnotes.
        footnotes_text = self._extract_footnotes(doc)
        if footnotes_text:
            metadata["footnotes"] = footnotes_text

        # If no title from properties, try the first heading.
        if not title and sections:
            title = sections[0].title

        # Extract references and figure/table refs from full text.
        references = self._extract_references_regex(full_text)
        fig_table_refs = self._detect_figure_table_refs(full_text)

        # Try to extract abstract.
        abstract = self._extract_abstract(full_text)

        return ParsedDocument(
            title=title,
            authors=authors,
            abstract=abstract,
            sections=sections,
            references=references,
            figures_tables=fig_table_refs,
            full_text=full_text,
            source_type="docx",
            source_path=source_path,
            metadata=metadata,
        )

    # ------------------------------------------------------------------
    # Section extraction
    # ------------------------------------------------------------------

    def _extract_sections(self, doc: Any) -> tuple[list[Section], str]:
        """Walk document paragraphs and group by heading styles.

        Returns:
            A tuple of ``(sections, full_text)``.
        """
        # section_data now includes the pre-computed character offset for each
        # section's start position in full_text, avoiding reverse-search with
        # find() which fails for synthetic headings like "Document".
        section_data: list[tuple[str, int, list[str], int]] = []
        current_heading = "Document"
        current_level = 1
        current_paragraphs: list[str] = []
        full_text_parts: list[str] = []
        # Track the character offset where the current section starts.
        # This is set when a section's first content paragraph is appended.
        current_section_start: int | None = None
        # Running character length of full_text as parts are appended.
        char_cursor = 0

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            text = para.text.strip()

            if not text:
                continue

            heading_level = _HEADING_STYLES.get(style_name)

            if heading_level is not None and heading_level > 0:
                # Save previous section.
                if current_paragraphs:
                    section_data.append(
                        (current_heading, current_level, current_paragraphs,
                         current_section_start if current_section_start is not None else char_cursor)
                    )
                current_heading = text
                current_level = heading_level
                current_paragraphs = []
                # The heading itself marks the start of the new section.
                # Record current char_cursor before appending.
                current_section_start = char_cursor
                # Account for the "\n\n" separator between parts.
                if full_text_parts:
                    char_cursor += 2  # len("\n\n")
                full_text_parts.append(text)
                char_cursor += len(text)
            elif heading_level == 0:
                # Title / Subtitle -- store but don't create a section.
                if full_text_parts:
                    char_cursor += 2
                full_text_parts.append(text)
                char_cursor += len(text)
            else:
                # Record offset of the first paragraph in this section.
                if not current_paragraphs and current_section_start is None:
                    current_section_start = char_cursor
                current_paragraphs.append(text)
                if full_text_parts:
                    char_cursor += 2
                full_text_parts.append(text)
                char_cursor += len(text)

        # Don't forget the last section.
        if current_paragraphs:
            section_data.append(
                (current_heading, current_level, current_paragraphs,
                 current_section_start if current_section_start is not None else char_cursor)
            )

        full_text = "\n\n".join(full_text_parts)

        # If no headings were found, wrap everything in one section.
        if not section_data:
            return (
                [
                    self._build_section(
                        section_id="sec-1",
                        title="Document",
                        text=full_text,
                        level=1,
                        start_char=0,
                    )
                ],
                full_text,
            )

        # Build Section objects using pre-computed character offsets.
        sections: list[Section] = []
        for idx, (heading, level, paragraphs, char_offset) in enumerate(section_data, start=1):
            body = "\n\n".join(paragraphs)
            section_id = f"sec-{idx}"
            section = self._build_section(
                section_id=section_id,
                title=heading,
                text=body,
                level=level,
                start_char=char_offset,
            )
            sections.append(section)

        return sections, full_text

    # ------------------------------------------------------------------
    # Table extraction
    # ------------------------------------------------------------------

    @staticmethod
    def _extract_tables(doc: Any) -> str:
        """Extract all tables from the document as plain text."""
        tables_parts: list[str] = []

        for idx, table in enumerate(doc.tables, start=1):
            rows_text: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows_text.append(" | ".join(cells))
            if rows_text:
                table_text = f"Table {idx}:\n" + "\n".join(rows_text)
                tables_parts.append(table_text)

        return "\n\n".join(tables_parts)

    # ------------------------------------------------------------------
    # Footnote extraction
    # ------------------------------------------------------------------

    @staticmethod
    def _extract_footnotes(doc: Any) -> str:
        """Extract footnotes from the DOCX XML.

        python-docx does not expose footnotes directly, so we parse
        the underlying XML when possible.
        """
        try:
            from docx.opc.constants import RELATIONSHIP_TYPE as RT  # type: ignore[import-untyped]

            footnotes_part = None
            for rel in doc.part.rels.values():
                if "footnotes" in rel.reltype:
                    footnotes_part = rel.target_part
                    break

            if footnotes_part is None:
                return ""

            from lxml import etree  # type: ignore[import-untyped]

            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            # Disable entity resolution to prevent XML bomb attacks (HIGH-S1).
            parser = etree.XMLParser(resolve_entities=False, no_network=True)
            tree = etree.fromstring(footnotes_part.blob, parser=parser)
            footnotes: list[str] = []
            for fn in tree.findall(".//w:footnote", ns):
                fn_id = fn.get(f"{{{ns['w']}}}id", "")
                # Skip separator footnotes (IDs 0 and -1).
                if fn_id in ("0", "-1"):
                    continue
                texts = [t.text or "" for t in fn.findall(".//w:t", ns)]
                fn_text = "".join(texts).strip()
                if fn_text:
                    footnotes.append(f"[{fn_id}] {fn_text}")

            return "\n".join(footnotes)
        except Exception:
            logger.debug("footnote_extraction_failed", exc_info=True)
            return ""

    # ------------------------------------------------------------------
    # Abstract extraction
    # ------------------------------------------------------------------

    @staticmethod
    def _extract_abstract(text: str) -> str | None:
        """Extract abstract from document text."""
        match = re.search(
            r"(?:^|\n)\s*Abstract[:\s]*\n?(.*?)(?=\n\s*(?:\d+\.?\s+)?(?:Introduction|Keywords|1\s)|$)",
            text,
            re.DOTALL | re.IGNORECASE,
        )
        if match:
            abstract = re.sub(r"\s+", " ", match.group(1).strip())
            if len(abstract) > 30:
                return abstract
        return None
